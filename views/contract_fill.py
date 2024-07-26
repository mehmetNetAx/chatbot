from docx import Document
import streamlit as st

def fill_document01(template_path, output_path, data):
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key,value)
                for run in paragraph.runs:
                    run.text = run.text.replace(key,value)
    
    doc.save(output_path)



def replace_text_in_paragraphs(paragraphs, data):
    for paragraph in paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key,value)
                for run in paragraph.runs:
                    run.text = run.text.replace(key, value)

def replace_text_in_tables(tables, data):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_paragraphs(cell.paragraphs, data)

def fill_document(template_path, output_path, data):
    doc = Document(template_path)
    
    # Replace keys in paragraphs
    replace_text_in_paragraphs(doc.paragraphs, data)
    
    # Replace keys in tables
    replace_text_in_tables(doc.tables, data)

    
    doc.save(output_path)




st.title("Papir.ai - Contract Generatior From Template")
tarih = st.date_input("Sözleşme Tarihi : ","today")
sirket_unvani = st.text_input("Şirket Ünvanı","NETAX BİLİŞİM HİZMETLERİ")
sirket_ticaret_unvani = st.text_input("Ticaret Ünvanı","NETAX BİLİŞİM ve DANIŞMANLIK HİZMETLERİ TİC. LTD. ŞTİ.")

sirket_adresi = st.text_input("Şirket Adresi","Altunizade mah. Dadaslar Sok. No:23/7 Üsküdar - İstanbul")
ticaret_sicilno = st.text_input("Ticaret Sicil No","45730-5")
firma_web_adresi = st.text_input("Firma Web Adresi","www.netaxtech.com")
firma_kep_adresi = st.text_input("Firma KEP Adresi","info@netaxtech.com")
firma_banka_adi = st.text_input("Firma Banka Adı","vakıf Bank")
firma_banka_sube = st.text_input("Firma Banka Şube","Kasımpaşa Şubesi")
firma_banka_hesapno = st.text_input("Firma Banka Hesap No","7307601354")
firma_banka_iban = st.text_input("Firma Banka IBAN","TR160001500158007307601354")
  
sozlesme_konusu = st.text_input("Sözleşme Konusu","Yazılım Geliştirme Hizmeti")

sozlesme_toplam_bedeli = st.number_input("Sözleşme Toplam Bedeli",min_value=300000)
sozlesme_bedel_madde1 = st.text_input("Sözleşme Bedeli Maddesi 1","%30 Peşin geri kalan eşit taksitle")

cezai_sart1 = st.text_input("Cezai Şart 1","Cezai Şart 1")
cezai_sart2 = st.text_input("Cezai Şart 2","Cezai Şart 2")
cezai_sart3 = st.text_input("Cezai Şart 3","Cezai Şart 3")

istenmeyen1 = st.text_input("esgis İstenmeyen Madde 1","İstenmeyen Madde 1")
istenmeyen2 = st.text_input("esgis İstenmeyen Madde 2","İstenmeyen Madde 2")
istenmeyen3 = st.text_input("esgis İstenmeyen Madde 3","İstenmeyen Madde 3")
istenmeyen4 = st.text_input("esgis İstenmeyen Madde 4","İstenmeyen Madde 4")

nusha_sayisi = st.number_input("Nüsha Sayısı",min_value=2)

sozlesme_suresi = st.number_input("Sözleşme Süresi - YIL",min_value=1)

fo_imza1 = st.text_input("FO İmza Yönetici 1","FO Yönetici 1")
fo_imza2 = st.text_input("FO İmza Yönetici 2","FO Yönetici 2")
fo_imza3 = st.text_input("FO İmza Yönetici 3","FO Yönetici 3")



firma_imza1 = st.text_input("Firma İmza 1","Firma İmza 1")
firma_imza2 = st.text_input("Firma İmza 2","Firma İmza 2")
firma_imza3 = st.text_input("Firma İmza 3","Firma İmza 3")

contract_type = st.selectbox("Type of Contract", ["Service", "Goods", "Employment", "Lease", "Other"])
startdate = st.date_input("Contract Start Date : ","today")
enddate = st.date_input("Contract End Date:","today")
termination_period = st.selectbox("Termination Period",["1 week Before","1 Month Before","Na"])
submit_button = st.button("Generate Contract")

def set_data():
        
    data = {
            '<tarih>' : f'{tarih}',
            '<Sirket Unvani>' : sirket_unvani,
            '<Sirket Adresi>' : sirket_adresi,
            '<ticaretsicilno>' : ticaret_sicilno,
            '<firmawebadresi>' : firma_web_adresi,
            '<firmaKEPadresi>' :  firma_kep_adresi,
            '<sozlesmenin_konusu>' : sozlesme_konusu,
            '<sozlesme_toplam_bedeli>' : f'{sozlesme_toplam_bedeli} TRL',
            '<sözleme_bedeli_maddesi_1>' : sozlesme_bedel_madde1,
            '<firma_banka_adi>' : firma_banka_adi,
            '<firma_banka_sube_adi>' : firma_banka_sube,
            '<firma_hesap_no>' : firma_banka_hesapno,
            '<firma_iban>' : firma_banka_iban,
            '<cezai_sart_1>' : cezai_sart1,
            '<cezai_sart_2>' : cezai_sart2,
            '<cezai_sart_3>' : cezai_sart3,
            
            '<esgis_istenmeyen_maddeler_1>' : istenmeyen1,
            '<esgis_istenmeyen_maddeler_2>' : istenmeyen2,
            '<esgis_istenmeyen_maddeler_3>' : istenmeyen3, 
            '<esgis_istenmeyen_maddeler_4>' : istenmeyen4,

            '<kac_nüsha>' : f'{nusha_sayisi} Nüsha',
            '<sozlesme_suresi>' : f' {sozlesme_suresi} Yıl ',
            '<firma_ticaret_unvani>' :  sirket_ticaret_unvani,
            '<FOimza1>' : fo_imza1,
            '<FOimza2>' : fo_imza2,
            '<FOimza3>' : fo_imza3,
            '<firmaimza1>' : firma_imza1,
            '<firmaimza2>' : firma_imza2,
            '<firmaimza3>' : firma_imza3

    }
    return data





if submit_button:
    st.subheader("Progressing...")
    data = set_data()
    template_path = 'FO_Contract_Template.docx'
    output_path = 'FO_Filled_Contract.docx'


    fill_document(template_path, output_path, data)
    
    st.subheader(f"{output_path} Sözleşmesi Yaratılmıştır...")
    
